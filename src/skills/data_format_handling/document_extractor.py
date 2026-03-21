"""
Document Extractor Skill

Extract text and metadata from PDF and DOCX documents.
"""
import json
import logging
import os
import sys
from typing import Any, Dict, List

import time
import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _get_pdfplumber():
    """Safely import pdfplumber."""
    try:
        import pdfplumber
        return pdfplumber
    except ImportError:
        return None


def _get_docx():
    """Safely import python-docx."""
    try:
        import docx
        return docx
    except ImportError:
        return None


def extract_pdf(file_path: str) -> Dict[str, Any]:
    """Extract text from PDF using pdfplumber."""
    pdf = _get_pdfplumber()
    if pdf is None:
        raise ImportError("pdfplumber not installed: pip install pdfplumber")
    
    result = {"pages": [], "metadata": {}, "text": ""}
    
    with pdf.open(file_path) as p:
        result["metadata"] = {
            "page_count": len(p.pages),
            "metadata": p.metadata
        }
        
        for i, page in enumerate(p.pages):
            page_text = page.extract_text() or ""
            result["pages"].append({"page": i + 1, "text": page_text})
            result["text"] += f"\n--- Page {i + 1} ---\n" + page_text
    
    return result


def extract_docx(file_path: str) -> Dict[str, Any]:
    """Extract text from DOCX using python-docx."""
    doc = _get_docx()
    if doc is None:
        raise ImportError("python-docx not installed: pip install python-docx")
    
    document = doc.Document(file_path)
    
    result = {"paragraphs": [], "tables": [], "metadata": {}}
    
    # Extract metadata
    core_props = document.core_properties
    result["metadata"] = {
        "author": core_props.author,
        "title": core_props.title,
        "subject": core_props.subject,
        "created": str(core_props.created) if core_props.created else None,
        "modified": str(core_props.modified) if core_props.modified else None,
    }
    
    # Extract paragraphs
    for para in document.paragraphs:
        if para.text.strip():
            result["paragraphs"].append({
                "text": para.text,
                "style": para.style.name if para.style else None
            })
    
    # Extract tables
    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        result["tables"].append(table_data)
    
    # Build full text
    result["text"] = "\n\n".join(p["text"] for p in result["paragraphs"])
    
    return result


async def invoke(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Main entry point for skill invocation."""
    start_time = time.time()
    timestamp = datetime.datetime.now().isoformat()
    
    action = payload.get("action", "extract")
    file_path = payload.get("file_path")
    output_path = payload.get("output_path")
    
    try:
        if action == "extract":
            if not file_path or not os.path.exists(file_path):
                return {"result": {"error": f"File not found: {file_path}"}, "metadata": {"timestamp": timestamp}}
            
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.pdf':
                data = extract_pdf(file_path)
            elif ext == '.docx':
                data = extract_docx(file_path)
            else:
                return {"result": {"error": f"Unsupported format: {ext}"}, "metadata": {"timestamp": timestamp}}
            
            # Optionally save to file
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2)
                data["output_file"] = output_path
            
            return {"result": data, "metadata": {"timestamp": timestamp, "format": ext}}
        
        elif action == "extract_text_only":
            if not file_path or not os.path.exists(file_path):
                return {"result": {"error": f"File not found: {file_path}"}, "metadata": {"timestamp": timestamp}}
            
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.pdf':
                data = extract_pdf(file_path)
            elif ext == '.docx':
                data = extract_docx(file_path)
            else:
                return {"result": {"error": f"Unsupported format: {ext}"}, "metadata": {"timestamp": timestamp}}
            
            return {"result": {"text": data["text"]}, "metadata": {"timestamp": timestamp}}
        
        elif action == "extract_metadata":
            if not file_path or not os.path.exists(file_path):
                return {"result": {"error": f"File not found: {file_path}"}, "metadata": {"timestamp": timestamp}}
            
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.pdf':
                data = extract_pdf(file_path)
            elif ext == '.docx':
                data = extract_docx(file_path)
            else:
                return {"result": {"error": f"Unsupported format: {ext}"}, "metadata": {"timestamp": timestamp}}
            
            return {"result": {"metadata": data["metadata"]}, "metadata": {"timestamp": timestamp}}
        
        else:
            return {"result": {"error": f"Unknown action: {action}", "available": ["extract", "extract_text_only", "extract_metadata"]}, "metadata": {"timestamp": timestamp}}
    
    except Exception as e:
        return {"result": {"error": str(e)}, "metadata": {"timestamp": timestamp, "error": str(e)}}


def register_skill():
    return {"name": "document-extractor", "description": "Extract text and metadata from PDF and DOCX documents", "version": "1.0.0", "domain": "DATA_FORMAT_HANDLING"}
